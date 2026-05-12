"""DOCX rapor üretici — python-docx ile FENlife template doldurma."""
import copy
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

_DEFAULT_TEMPLATE_PATH = Path(__file__).parent / "templates" / "fenlife_base.docx"

SUBJECT_KEYS = ["turkce", "matematik", "fen", "sosyal", "din", "ydil"]
SUBJECT_LABELS = {
    "turkce": "Türkçe",
    "matematik": "Matematik",
    "fen": "Fen Bilimleri",
    "sosyal": "İnkılap",
    "din": "Din Kültürü",
    "ydil": "Yabancı Dil",
}

FENLIFE_BLUE = RGBColor(0x1B, 0x5E, 0x8C)

_TR_UPPER = str.maketrans("çğıiöşü", "ÇĞIİÖŞÜ")


def _tr_upper(s: str) -> str:
    """Türkçe karakterleri doğru büyülten uppercase."""
    return s.translate(_TR_UPPER).upper()


# ---------------------------------------------------------------------------
# Cell helpers
# ---------------------------------------------------------------------------

def _cell_set_text(cell, text: str, bold: bool = False, color: RGBColor | None = None):
    """Clear all cell paragraphs, write text into first paragraph's first run."""
    # Remove extra paragraphs beyond the first (old template placeholders)
    for extra_para in list(cell.paragraphs[1:]):
        p = extra_para._p
        p.getparent().remove(p)
    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ""
    run = para.runs[0] if para.runs else para.add_run()
    run.text = text
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def _find_and_replace_in_cell(cell, old: str, new: str):
    """Replace placeholder text within a cell, handling split-run placeholders."""
    # Rebuild full paragraph text and replace
    for para in cell.paragraphs:
        full = "".join(r.text for r in para.runs)
        if old not in full:
            continue
        replaced = full.replace(old, new)
        # Put replaced text into first run, clear others
        if para.runs:
            para.runs[0].text = replaced
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.add_run(replaced)


def _replace_in_paragraph(para, old: str, new: str):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    replaced = full.replace(old, new)
    if para.runs:
        para.runs[0].text = replaced
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(replaced)


def _copy_row(table, source_row_idx: int):
    """Append a deep copy of source_row to table, return the new row."""
    source_row = table.rows[source_row_idx]
    new_tr = copy.deepcopy(source_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _delete_row(table, row_idx: int):
    tr = table.rows[row_idx]._tr
    table._tbl.remove(tr)


def _write_summary_to_cell(cell, text: str) -> None:
    """Risk değerlendirme hücresine özet metni yaz (ikinci paragrafa direkt yaz, fallback: içerik ara)."""
    paras = cell.paragraphs
    if len(paras) > 1:
        para = paras[1]
        if para.runs:
            para.runs[0].text = text
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.add_run(text)
    else:
        # Tek paragraf varsa kısmi placeholder ile dene
        _find_and_replace_in_cell(cell, "[Öğ", text)


def _fill_action_table(table, actions: list[tuple], max_rows: int = 3):
    """
    Aksiyon tablosunu (rehberlik / teacher) doldur.
    Header satırı (row 0) korunur, veri satırları rows 1..max_rows.
    Fazla / doldurulamayan satırlar silinir.
    """
    for ai, (act_title, act_detail) in enumerate(actions[:max_rows]):
        row_idx = ai + 1
        if row_idx >= len(table.rows):
            break
        cell = table.rows[row_idx].cells[0]
        paras = cell.paragraphs
        if paras:
            _replace_in_paragraph(paras[0], paras[0].text.strip(), f"▸  {act_title}")
        if len(paras) > 1:
            _replace_in_paragraph(paras[1], paras[1].text.strip(), act_detail)

    # Doldurulamayan satırları temizle (placeholder metni kaldır)
    filled = len(actions)
    for row_idx in range(filled + 1, len(table.rows)):
        cell = table.rows[row_idx].cells[0]
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ""


# ---------------------------------------------------------------------------
# Algorithmic text builders
# ---------------------------------------------------------------------------

RISK_TITLES = {
    "dusuk": "Stabil Gelişim",
    "orta": "Orta Öncelikli",
    "yuksek": "Yüksek Öncelikli",
    "sinirli": "Sınırlı Analiz",
}


def _subject_comment(label: str, avg: float, pct: float) -> str:
    if pct >= 75:
        return f"Güçlü · sürdürülebilir"
    if pct >= 55:
        return f"Orta · hedefli çalışma"
    if pct >= 35:
        return f"Zayıf · öncelikli konu tekrarı"
    return "KRİTİK · acil müdahale gerekli"


def _build_findings(subject_avgs: list[dict], trend: dict, risk: dict,
                    worst_gap: dict | None, exam_count: int = 0,
                    avg_puan: float = 0.0) -> list[tuple[str, str]]:
    """Return up to 4 (title, detail) pairs for the Kritik Bulgular section."""
    findings = []

    if not subject_avgs:
        findings.append(("Ders verisi eksik", "Bu rapor için ders bazlı net verisi bulunmadığından detaylı analiz üretilemedi."))
        findings.append(("Veri yetersiz", "Daha kapsamlı analiz için tüm derslerden net verisi içeren sınav sonuçları yükleyin."))
        return findings

    # 1. Weakest subject
    weakest = min(subject_avgs, key=lambda s: s["pct"])
    findings.append((
        f"{weakest['label']} kritik zayıflık",
        f"Ortalama {weakest['avg']:.1f} net, başarı oranı %{weakest['pct']:.0f}. "
        f"En düşük performanslı ders — öncelikli müdahale gerekiyor.",
    ))

    # 2. Trend (sınırlı veriye uyarlanır)
    direction = trend.get("direction", "flat")
    delta = trend.get("delta", 0)
    if exam_count < 2:
        findings.append((
            "Tek sınav verisi — anlık görüntü",
            f"Eğilim analizi için en az 2 deneme gerekir. "
            f"Bu sınavdaki puan: {avg_puan:.0f}. "
            f"Güvenilir eğilim ve karşılaştırmalı analiz için en az 5 sınav yükleyin.",
        ))
    elif direction == "up":
        t_title = "Yükselen eğilim"
        t_desc = f"Son denemeden ilk denemeye kıyasla +{abs(delta):.0f} puan artış. Motivasyon ve çalışma düzeni desteklenebilir."
        findings.append((t_title, t_desc))
    elif direction == "down":
        t_title = "Düşen eğilim — dikkat"
        t_desc = f"Son denemede {abs(delta):.0f} puanlık gerileme. Sınav kaygısı veya çalışma sürekliliği sorgulanabilir."
        findings.append((t_title, t_desc))
    else:
        t_title = "Sabit seyir"
        t_desc = f"Puan serisi düz seyrediyor (delta: {delta:.0f}). Yeni kazanım pratiği platoya çıkışı tetikleyebilir."
        findings.append((t_title, t_desc))

    # 3. Difficulty gap veya güçlü ders
    if worst_gap and worst_gap["gap"] > 10:
        findings.append((
            f"{worst_gap['subject']}: {worst_gap['level']} seviye açığı",
            f"Kurum ortalaması %{worst_gap['general_pct']:.0f}, öğrenci %{worst_gap['student_pct']:.0f} — "
            f"{worst_gap['gap']:.0f} puan fark. Sınav stratejisi ve dikkat yönetimi çalışması önerilir.",
        ))
    else:
        strongest = max(subject_avgs, key=lambda s: s["pct"])
        findings.append((
            f"{strongest['label']} güçlü yan",
            f"Başarı oranı %{strongest['pct']:.0f} ile en yüksek performanslı ders. "
            f"Zor soru pratiğiyle puan katkısı artırılabilir.",
        ))

    # 4. Risk assessment
    r_label = RISK_TITLES.get(risk.get("level", "orta"), "Orta Öncelikli")
    findings.append((
        f"Genel risk seviyesi: {r_label}",
        f"Risk sınıflandırması: ortalama puan, trend yönü ve son deneme sapması birlikte değerlendirilerek belirlendi.",
    ))

    return findings[:4]


def _build_teacher_actions(subject_avgs: list[dict], priority_topics: list[dict]) -> list[dict]:
    """Build up to 3 branch-teacher action blocks."""
    if not subject_avgs:
        return []
    sorted_subj = sorted(subject_avgs, key=lambda s: s["pct"])
    top3_subj = sorted_subj[:3]
    blocks = []
    for subj in top3_subj:
        related = [t for t in priority_topics if t["subject"] == subj["label"]][:2]
        actions = [
            (
                f"Konu tekrar planı — {subj['label']}",
                f"Ortalama {subj['avg']:.1f} net, başarı %{subj['pct']:.0f}. "
                f"Zayıf konulara odaklı haftalık çalışma kağıdı düzenlenebilir.",
            ),
        ]
        for t in related:
            actions.append((
                f"Kazanım: {t['topic'][:50]}",
                f"{t['q']} soruda %{t['success']:.0f} başarı — mini çözümlü anlatım + 5-8 soruluk alıştırma önerilir.",
            ))
        blocks.append({"subject": subj["label"], "actions": actions[:3]})
    return blocks


def _build_guidance_actions(student_data: dict, risk: dict, trend: dict,
                             subject_avgs: list[dict]) -> tuple[list, list]:
    """Return (strong_side_actions, family_actions)."""
    if not subject_avgs:
        return [], [("Veri yetersiz", "Rehberlik önerileri için ders verisi gereklidir.")]
    strongest = max(subject_avgs, key=lambda s: s["pct"])
    strong_actions = [
        (
            f"{strongest['label']} güçlü yan olarak kullan",
            f"Başarı %{strongest['pct']:.0f}. Zor soru pratiği ekleyerek özgüven inşasını destekle.",
        ),
        (
            "Hedef belirleme seansı",
            f"Mevcut ortalama {student_data.get('avg_puan', 0):.0f} puan üzerinden somut LGS hedefi netleştir.",
        ),
    ]
    family = [
        (
            "Çalışma düzeni paylaşımı",
            "Haftalık ders çalışma saatleri ve tekrar sıklığı aileyle paylaşılabilir.",
        ),
    ]
    if risk.get("level") == "yuksek" or trend.get("direction") == "down":
        family.append((
            "Motivasyon görüşmesi",
            f"Düşen trend ({trend.get('delta', 0):.0f} delta) ve yüksek risk seviyesi. "
            f"Sınav kaygısı ve motivasyon aile görüşmesine taşınabilir.",
        ))
    return strong_actions[:2], family[:2]


def _build_follow_up_targets(weakest: dict, avg_puan: float, priority_topics: list[dict],
                              next_exam_date: str) -> list[tuple]:
    """5 rows for follow-up targets table (excluding header)."""
    return [
        (
            f"{weakest['label']} — net",
            f"{weakest['avg']:.1f}",
            f"{weakest['avg'] + 2:.1f}",
            "+2 net artış",
        ),
        (
            "Genel LGS puanı",
            f"{avg_puan:.0f}",
            f"{avg_puan + 15:.0f}",
            "+15 puan artış",
        ),
        (
            f"Öncelikli kazanım başarısı",
            f"%{priority_topics[0]['success']:.0f}" if priority_topics else "—",
            "%50+",
            "Mini test ile ölçüm",
        ),
        (
            "Tekrar eden hata azaltımı",
            "Mevcut",
            "−1 yanlış/deneme",
            "Hata analiz defteri",
        ),
        (
            "Haftalık çalışma saati",
            "Veri yok",
            "12 saat",
            "Çalışma takip defteri",
        ),
    ]


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------

class ReportGenerator:
    FENLIFE_BLUE = "1B5E8C"
    FENLIFE_ORANGE = "E87722"

    def generate(
        self,
        student_data: dict,
        charts_dir: Path | None,
        output_path: Path,
        template_path: Path | None = None,
    ) -> Path:
        """
        FENlife template'ini veriyle doldur.

        student_data keys:
            name, grade, avg_puan, risk, trend, exams, subject_avgs,
            priority_topics, difficulty, worst_gap

        template_path: override for REPORT_TEMPLATE_PATH config; falls back to
            _DEFAULT_TEMPLATE_PATH if None or missing.
        """
        from app.config import settings

        resolved_template = (
            template_path
            or (settings.report_template_path if settings.report_template_path.exists() else None)
            or _DEFAULT_TEMPLATE_PATH
        )
        doc = Document(resolved_template)
        tables = doc.tables

        name = student_data["name"]
        grade = student_data.get("grade", "8-A")
        avg_puan = student_data["avg_puan"]
        risk = student_data["risk"]
        trend = student_data["trend"]
        exams = student_data["exams"]
        subject_avgs = student_data["subject_avgs"]
        priority_topics = student_data.get("priority_topics", [])
        difficulty = student_data.get("difficulty", {})
        worst_gap = student_data.get("worst_gap")
        today = date.today().strftime("%d.%m.%Y")
        publishers = list({e.get("publisher", "") for e in exams if e.get("publisher")})
        pub_str = ", ".join(publishers) if publishers else "—"

        # ----- Table 1: student info bar -----
        t1 = tables[1]
        _cell_set_text(t1.rows[0].cells[0], f"ÖĞRENCİ\n{name}", bold=False)
        _cell_set_text(t1.rows[0].cells[1], f"SINIF\n{grade}")
        _cell_set_text(t1.rows[0].cells[2], f"DENEME SAYISI\n{len(exams)}")
        _cell_set_text(t1.rows[0].cells[3], f"ORTALAMA PUAN\n{avg_puan:.1f}")

        # ----- Table 2: risk evaluation -----
        risk_title = RISK_TITLES.get(risk.get("level", "orta"), "Orta Öncelikli")
        t2 = tables[2]
        t2_cell = t2.rows[0].cells[0]
        _find_and_replace_in_cell(t2_cell, "[Genel Başlık — ör. Stabil Gelişim]", risk_title)
        analysis_tier = student_data.get("analysis_tier", {})
        if len(exams) == 1:
            summary = (
                f"Tek sınav verisiyle anlık görüntü analizi yapıldı. "
                f"Bu sınavdaki puan: {avg_puan:.1f}. "
                f"Kesin eğilim ve gelişim karşılaştırması için en az 5 sınav verisi gereklidir."
            )
        else:
            trend_label = (
                "yükseliş" if trend.get("direction") == "up"
                else "düşüş" if trend.get("direction") == "down"
                else "sabit seyir"
            )
            summary = (
                f"{len(exams)} deneme boyunca ders bazlı net ve puan hareketi incelendi. "
                f"Ortalama puan {avg_puan:.1f}, trend yönü: {trend_label}."
            )
        if analysis_tier.get("warning"):
            summary += f" ⚠ {analysis_tier['warning']}"
        _write_summary_to_cell(t2_cell, summary)

        # ----- Paragraphs: summary -----
        for para in doc.paragraphs:
            _replace_in_paragraph(para, "[N] deneme boyunca ders bazlı net ve puan hareketi.", summary)
            _replace_in_paragraph(para, "[Anomali varsa burada not düşülür — ör. Mozaik 4 (08.03.2026) anomali olarak işaretlendi]", "")
            _replace_in_paragraph(para, "[TARİH]", today)
            _replace_in_paragraph(para, "[Yayın Adı]", pub_str)

        # ----- Table 3: exam rows -----
        t3 = tables[3]
        # Row 0 = header, rows 1..N-2 = data template rows, last row = Ortalama
        # Save template data row and avg row, delete template rows
        avg_row_tr = copy.deepcopy(t3.rows[-1]._tr)
        template_data_tr = copy.deepcopy(t3.rows[1]._tr)

        # Delete all rows except header (row 0)
        while len(t3.rows) > 1:
            _delete_row(t3, len(t3.rows) - 1)

        # Add exam rows
        for exam in exams:
            new_tr = copy.deepcopy(template_data_tr)
            t3._tbl.append(new_tr)
            row = t3.rows[-1]
            cells = row.cells
            cells[0].paragraphs[0].runs[0].text = exam.get("date", "")[-5:].replace("-", ".")  # MM.DD
            cells[1].paragraphs[0].runs[0].text = exam.get("name", "")
            cells[2].paragraphs[0].runs[0].text = exam.get("publisher", "")
            for i, key in enumerate(SUBJECT_KEYS):
                val = exam.get(key, 0)
                cells[3 + i].paragraphs[0].runs[0].text = f"{val:.2f}"
            cells[9].paragraphs[0].runs[0].text = f"{exam.get('puan', 0):.0f}"

        # Add averages row
        t3._tbl.append(avg_row_tr)
        avg_row = t3.rows[-1]
        avg_cells = avg_row.cells
        avg_cells[0].paragraphs[0].runs[0].text = "Ortalama"
        # Clear merged cells
        for ci in [1, 2]:
            if avg_cells[ci].paragraphs[0].runs:
                avg_cells[ci].paragraphs[0].runs[0].text = ""
        for i, sa in enumerate(subject_avgs):
            if 3 + i < len(avg_cells):
                if avg_cells[3 + i].paragraphs[0].runs:
                    avg_cells[3 + i].paragraphs[0].runs[0].text = f"{sa['avg']:.2f}"
                else:
                    avg_cells[3 + i].paragraphs[0].add_run(f"{sa['avg']:.2f}")
        if avg_cells[9].paragraphs[0].runs:
            avg_cells[9].paragraphs[0].runs[0].text = f"{avg_puan:.0f}"

        # ----- Tables 4-7: Kritik Bulgular -----
        findings = _build_findings(
            subject_avgs, trend, risk, worst_gap,
            exam_count=len(exams), avg_puan=avg_puan,
        )
        for fi, (title, detail) in enumerate(findings):
            cell = tables[4 + fi].rows[0].cells[0]
            full = "".join(r.text for p in cell.paragraphs for r in p.runs)
            # Replace placeholder title and detail in the two paragraphs
            paras = cell.paragraphs
            if len(paras) >= 1:
                _replace_in_paragraph(paras[0], paras[0].text.strip(), f"{fi + 1}.  {title}")
            if len(paras) >= 2:
                _replace_in_paragraph(paras[1], paras[1].text.strip(), detail)

        # ----- Table 8: subject averages -----
        t8 = tables[8]
        for i, sa in enumerate(subject_avgs[:6]):
            row = t8.rows[i + 1]
            cells = row.cells
            _cell_set_text(cells[0], sa["label"])
            _cell_set_text(cells[1], f"{sa['avg']:.2f}")
            _cell_set_text(cells[2], f"{sa.get('max', sa['avg']):.2f}")
            _cell_set_text(cells[3], f"{sa.get('min', sa['avg']):.2f}")
            _cell_set_text(cells[4], f"%{sa['pct']:.0f}")
            _cell_set_text(cells[5], _subject_comment(sa["label"], sa["avg"], sa["pct"]))

        # ----- Table 9: priority topics -----
        t9 = tables[9]
        template_topic_tr = copy.deepcopy(t9.rows[1]._tr)
        # Delete template rows (keep header)
        while len(t9.rows) > 1:
            _delete_row(t9, len(t9.rows) - 1)
        for idx, topic in enumerate(priority_topics[:10]):
            new_tr = copy.deepcopy(template_topic_tr)
            t9._tbl.append(new_tr)
            row = t9.rows[-1]
            cells = row.cells
            topic_text = topic.get("topic", "")
            if topic.get("estimated"):
                topic_text = f"(Tahmini) {topic_text}"
            q_val = topic.get("q", 0)
            q_display = "—" if topic.get("estimated") else str(q_val)
            if cells[0].paragraphs[0].runs:
                cells[0].paragraphs[0].runs[0].text = str(idx + 1)
            if cells[1].paragraphs[0].runs:
                cells[1].paragraphs[0].runs[0].text = topic.get("subject", "")
            if cells[2].paragraphs[0].runs:
                cells[2].paragraphs[0].runs[0].text = topic_text
            if cells[3].paragraphs[0].runs:
                cells[3].paragraphs[0].runs[0].text = q_display
            if cells[4].paragraphs[0].runs:
                cells[4].paragraphs[0].runs[0].text = f"%{topic.get('success', 0):.0f}"
            if cells[5].paragraphs[0].runs:
                cells[5].paragraphs[0].runs[0].text = f"%{topic.get('priority', 0):.0f}"

        # ----- Tables 10-12: Branch teacher actions -----
        teacher_blocks = _build_teacher_actions(subject_avgs, priority_topics)
        for bi, block in enumerate(teacher_blocks[:3]):
            t = tables[10 + bi]
            # Row 0: subject header
            if t.rows[0].cells[0].paragraphs[0].runs:
                t.rows[0].cells[0].paragraphs[0].runs[0].text = f"◆  {_tr_upper(block['subject'])} — Öncelik {bi + 1}"
            # Action rows: rows 1..
            for ai, (act_title, act_detail) in enumerate(block["actions"]):
                if ai + 1 < len(t.rows):
                    cell = t.rows[ai + 1].cells[0]
                    paras = cell.paragraphs
                    if paras:
                        _replace_in_paragraph(paras[0], paras[0].text.strip(), f"▸  {act_title}")
                    if len(paras) > 1:
                        _replace_in_paragraph(paras[1], paras[1].text.strip(), act_detail)
                    elif paras:
                        paras[0].add_run(f"\n{act_detail}")

        # ----- Tables 13-14: Guidance actions -----
        strong_actions, family_actions = _build_guidance_actions(
            student_data, risk, trend, subject_avgs
        )
        _fill_action_table(tables[13], strong_actions, max_rows=2)
        _fill_action_table(tables[14], family_actions, max_rows=2)

        # ----- Table 15: Study material inputs -----
        t15 = tables[15]
        template_study_tr = copy.deepcopy(t15.rows[1]._tr)
        while len(t15.rows) > 1:
            _delete_row(t15, len(t15.rows) - 1)
        for topic in priority_topics[:5]:
            new_tr = copy.deepcopy(template_study_tr)
            t15._tbl.append(new_tr)
            row = t15.rows[-1]
            cells = row.cells
            _cell_set_text(cells[0], topic.get("topic", "")[:60])
            _cell_set_text(cells[1], topic.get("subject", ""))
            q = topic.get("q", 0)
            _cell_set_text(cells[2], str(min(q + 2, 8)))
            # Difficulty distribution suggestion: more kolay if success low
            s = topic.get("success", 50)
            if s < 40:
                dist = "3 kolay · 3 orta · 2 zor"
            elif s < 65:
                dist = "2 kolay · 3 orta · 3 zor"
            else:
                dist = "1 kolay · 2 orta · 4 zor"
            _cell_set_text(cells[3], dist)

        # ----- Table 16: Follow-up targets -----
        t16 = tables[16]
        if subject_avgs:
            weakest = min(subject_avgs, key=lambda s: s["pct"])
            targets = _build_follow_up_targets(weakest, avg_puan, priority_topics, today)
            for i, (alan, mevcut, hedef, olcut) in enumerate(targets):
                if i + 1 < len(t16.rows):
                    cells = t16.rows[i + 1].cells
                    _cell_set_text(cells[0], alan)
                    _cell_set_text(cells[1], mevcut)
                    _cell_set_text(cells[2], hedef)
                    _cell_set_text(cells[3], olcut)

        # ----- Insert charts if available -----
        if charts_dir:
            _insert_chart_after(doc, "puan_trend.png",   charts_dir, anchor="01  Öğrenci Özeti",     width=5.5)
            _insert_chart_after(doc, "ders_basari.png",  charts_dir, anchor="03  Ders Bazlı Analiz",  width=4.5)
            _insert_chart_after(doc, "kazanimlar.png",   charts_dir, anchor="Öncelikli Kazanımlar",   width=5.0)
            _insert_chart_after(doc, "zorluk_matrix.png",charts_dir, anchor="[Kritik Ders] Derinlemesine", width=5.0)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path


def _insert_chart_after(doc: Document, filename: str, charts_dir: Path, anchor: str, width: float = 5.5):
    """Insert chart PNG as inline image after the first paragraph whose text contains anchor."""
    chart_path = charts_dir / filename
    if not chart_path.exists():
        return
    for para in doc.paragraphs:
        if anchor in para.text:
            run = para.add_run()
            run.add_picture(str(chart_path), width=Inches(width))
            break


_generator = ReportGenerator()


def generate_report(
    student_data: dict,
    output_path: Path,
    charts_dir: Path | None = None,
    template_path: Path | None = None,
) -> Path:
    """Module-level convenience wrapper."""
    return _generator.generate(student_data, charts_dir, output_path, template_path)
